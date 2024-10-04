

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import RGBColor
import requests
def convert_json_to_doc_buffer(response:dict) -> BytesIO:
    doc = Document()
    # black heading color 
    def add_heading_with_color(text, level):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  
            run.font.name = 'Times New Roman'  
            run.font.size = Pt(12)  
        doc.add_paragraph() 

    # Title
    title_heading = doc.add_heading(response["title"]["text"].upper(), level=1)
    title_run = title_heading.runs[0]
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
    counter = 1

    # Background atteched with technical_field
    add_heading_with_color('BACKGROUND', level=1)
    
    # Add technical field text with numbering at the start
    technical_field = response["technical_field"]["text"]
    technical_paragraph = doc.add_paragraph()
    technical_run = technical_paragraph.add_run(f"[{counter:04d}] ")
    technical_run.bold = True
    technical_run.font.name = 'Times New Roman'
    technical_run.font.size = Pt(12)
    tech_field_run = technical_paragraph.add_run(technical_field)
    tech_field_run.font.name = 'Times New Roman'
    tech_field_run.font.size = Pt(12)
    technical_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1  
    # background text 
    background_text = response["background"]["text"]
    sections = background_text.split('\n\n')
    for section in sections:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[{counter:04d}] ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        section_run = paragraph.add_run(section)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        counter += 1   
    
    #Breif summary text 
    add_heading_with_color('BRIEF SUMMARY', level=1)  
    summary_field = response["summary"]["text"]
    summary_paragraph = doc.add_paragraph()
    summary_run = summary_paragraph.add_run(f"[{counter:04d}] ")
    summary_run.bold = True
    summary_run.font.name = 'Times New Roman'
    summary_run.font.size = Pt(12)
    summary_field_run = summary_paragraph.add_run(summary_field)
    summary_field_run.font.name = 'Times New Roman'
    summary_field_run.font.size = Pt(12)
    summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1  

    # List of figures 
    add_heading_with_color('BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS', level=1)
    for figure in response["list_of_figures"]:
        figure_paragraph = doc.add_paragraph()
        figure_run = figure_paragraph.add_run(f"[{counter:04d}] ")
        figure_run.bold = True
        figure_run.font.name = 'Times New Roman'
        figure_run.font.size = Pt(12)
        figure_text_run = figure_paragraph.add_run(figure)
        figure_text_run.font.name = 'Times New Roman'
        figure_text_run.font.size = Pt(12)
        figure_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        counter += 1  
   
    #Detailed Description 
    add_heading_with_color('DETAILED DESCRIPTION', level=1)
    #method description
    method_desc_text = response["description"]["method_desc"]["text"]
    method_desc_sections = method_desc_text.split('\n\n')

    for section in method_desc_sections:
        method_desc_paragraph = doc.add_paragraph()
        method_desc_run = method_desc_paragraph.add_run(f"[{counter:04d}] ")
        method_desc_run.bold = True
        method_desc_run.font.name = 'Times New Roman'
        method_desc_run.font.size = Pt(12)
        section_run = method_desc_paragraph.add_run(section)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        method_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        counter += 1   
        
    #system description
    system_desc_text_list = response["description"]["system_desc"]["text_list"]
    for index, section in enumerate(system_desc_text_list):
        system_desc_paragraph = doc.add_paragraph()
        system_desc_run = system_desc_paragraph.add_run(f"[{counter:04d}] ")
        system_desc_run.bold = True
        system_desc_run.font.name = 'Times New Roman'
        system_desc_run.font.size = Pt(12)
        section_run = system_desc_paragraph.add_run(section)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        system_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        counter += 1 
    
    # invention Description 
    invention_desc_text = response["description"]["invention_desc"]["text"]
    invention_desc_sections = invention_desc_text.split('\n\n')
    for section in invention_desc_sections:
        invention_desc_paragraph = doc.add_paragraph()
        invention_desc_run = invention_desc_paragraph.add_run(f"[{counter:04d}] ")
        invention_desc_run.bold = True
        invention_desc_run.font.name = 'Times New Roman'
        invention_desc_run.font.size = Pt(12)
        section_run = invention_desc_paragraph.add_run(section)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        invention_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        counter += 1  
    
    # Claims Description  section 
    doc.add_page_break() 
    claims_heading = doc.add_heading('CLAIMS', level=1)
    claims_run = claims_heading.runs[0]
    claims_run.font.color.rgb = RGBColor(0, 0, 0)   
    claims_run.font.name = 'Times New Roman'
    claims_run.font.size = Pt(12)
    claims_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  

    doc.add_paragraph("What is claimed is:")
    claim_indent = Pt(20)
    for index, claim in enumerate(response['claims'], start=1):
            claim_paragraph = doc.add_paragraph()
            claim_paragraph.paragraph_format.left_indent = claim_indent

            if claim['tag'] == 'dependent':
                # Add spaces and numbering for dependent claims
                claim_run = claim_paragraph.add_run(f"{index}. ")   
                claim_paragraph.paragraph_format.left_indent = Pt(40)  
            else:
                claim_run = claim_paragraph.add_run(f"{index}. ")

            # Set font for the numbering
            claim_run.font.name = 'Times New Roman'
            claim_run.font.size = Pt(12)
            
            claim_text_run = claim_paragraph.add_run(claim['text'])
            claim_text_run.font.name = 'Times New Roman'
            claim_text_run.font.size = Pt(12)
            claim_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Abstract text 
    doc.add_page_break() 
    abstract_heading = doc.add_heading('ABSTRACT', level=1)
    doc.add_paragraph()
    abstract_run = abstract_heading.runs[0]
    abstract_run.font.color.rgb = RGBColor(0, 0, 0)   
    abstract_run.font.name = 'Times New Roman'
    abstract_run.font.size = Pt(12)
    abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
    abstract_paragraph = doc.add_paragraph(response["abstract"]["text"])
    for run in abstract_paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    abstract_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    abstract_paragraph.paragraph_format.first_line_indent = Pt(50)

    # Figures
    doc.add_page_break() 
    add_heading_with_color('FIGURES', level=1)
    sorted_claims = sorted(response["claims"], key=lambda x: x.get("claim_type") == "system")
    for claim in sorted_claims:
     if claim.get("generated_figures_data") and claim["generated_figures_data"].get("latex_details"):
        claim_type = claim.get("claim_type")
        for latex_detail in claim["generated_figures_data"]["latex_details"]:
            if "images_urls" in latex_detail:
                for img_url in latex_detail["images_urls"]:
                    # print(f"this is ClaimType.{claim_type.upper()} url \n", img_url)
                    img_response = requests.get(img_url)
                    if img_response.status_code == 200:
                        image_stream = BytesIO(img_response.content)
                        image_paragraph = doc.add_paragraph()
                        run = image_paragraph.add_run()
                        run.add_picture(image_stream, width=Inches(4))
                        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    else:
                        print(f"Failed to download the image from {img_url}")
 
    # Add "Attorney Docket No." to the header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_paragraph.add_run("Attorney Docket No.")

    # Function to add page numbering in footer 
    def add_page_numbering(doc):
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.add_paragraph()
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = footer_paragraph.add_run()
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
        )
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE  \\* MERGEFORMAT </w:instrText>'
            )
        )
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
        )

    add_page_numbering(doc)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  
    return buffer





 # this is original code for claims 
    # doc.add_paragraph("What is claimed is:")
    # claim_indent = Pt(20)
    # for index, claim in enumerate(response['claims'], start=1):
    #     claim_paragraph = doc.add_paragraph()
    #     claim_paragraph.paragraph_format.left_indent = claim_indent

    #     claim_run = claim_paragraph.add_run(f"{index}. ")
    #     if claim['tag'] == 'dependent':
    #         claim_paragraph.paragraph_format.left_indent = Pt(
    #             40)  # Indent for dependent claims

    #     claim_run.font.name = 'Times New Roman'
    #     claim_run.font.size = Pt(12)
    #     claim_text_run = claim_paragraph.add_run(claim['text'])
    #     claim_text_run.font.name = 'Times New Roman'
    #     claim_text_run.font.size = Pt(12)
    #     claim_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # here another cliams start to show the data 
 





#development code 4october 
def convert_json_to_doc_buffer(response: dict) -> BytesIO:
    doc = Document()

    def add_heading_with_color(text, level):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        doc.add_paragraph()

    def add_paragraph_with_numbering(text, counter):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[{counter:04d}] ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        section_run = paragraph.add_run(text)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        return counter + 1  # Increment counter

    # Title
    title_heading = doc.add_heading(response["title"]["text"].upper(), level=1)
    title_run = title_heading.runs[0]
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    counter = 1

    # Background attached with technical field
    add_heading_with_color('BACKGROUND', level=1)

    # Add technical field text with numbering
    counter = add_paragraph_with_numbering(
        response["technical_field"]["text"], counter)

    # Background text
    background_text = response["background"]["text"]
    sections = background_text.split('\n\n')
    for section in sections:
        counter = add_paragraph_with_numbering(section, counter)

    # Brief summary text
    add_heading_with_color('BRIEF SUMMARY', level=1)
    counter = add_paragraph_with_numbering(
        response["summary"]["text"], counter)

    # List of figures
    add_heading_with_color(
        'BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS', level=1)
    for figure in response["list_of_figures"]:
        counter = add_paragraph_with_numbering(figure, counter)

    # Detailed Description
    add_heading_with_color('DETAILED DESCRIPTION', level=1)

    # Method description
    method_desc_text = response["description"]["method_desc"]["text"]
    method_desc_sections = method_desc_text.split('\n\n')
    for section in method_desc_sections:
        counter = add_paragraph_with_numbering(section, counter)

    # System description
    system_desc_text_list = response["description"]["system_desc"]["text_list"]
    for section in system_desc_text_list:
        counter = add_paragraph_with_numbering(section, counter)

    # Invention Description
    invention_desc_text = response["description"]["invention_desc"]["text"]
    invention_desc_sections = invention_desc_text.split('\n\n')
    for section in invention_desc_sections:
        counter = add_paragraph_with_numbering(section, counter)

    # Claims Description section
    doc.add_page_break()
    claims_heading = doc.add_heading('CLAIMS', level=1)
    claims_run = claims_heading.runs[0]
    claims_run.font.color.rgb = RGBColor(0, 0, 0)
    claims_run.font.name = 'Times New Roman'
    claims_run.font.size = Pt(12)
    claims_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Claims
    doc.add_paragraph("What is claimed is:")
    doc.add_paragraph()
    claim_indent = Pt(40)
    for index, claim in enumerate(response['claims'], start=1):
        claim_paragraph = doc.add_paragraph()
        claim_paragraph.paragraph_format.first_line_indent = claim_indent 
        
        claim_run = claim_paragraph.add_run(f"{index}. ")
        claim_run.font.name = 'Times New Roman'
        claim_run.font.size = Pt(12)
        
        claim_text_run = claim_paragraph.add_run(claim['text'])
        claim_text_run.font.name = 'Times New Roman'
        claim_text_run.font.size = Pt(12)
        doc.add_paragraph()
         
    
    # Abstract text
    doc.add_page_break()
    abstract_heading = doc.add_heading('ABSTRACT', level=1)
    doc.add_paragraph()
    abstract_run = abstract_heading.runs[0]
    abstract_run.font.color.rgb = RGBColor(0, 0, 0)
    abstract_run.font.name = 'Times New Roman'
    abstract_run.font.size = Pt(12)
    abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    abstract_paragraph = doc.add_paragraph(response["abstract"]["text"])

    for run in abstract_paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    abstract_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    abstract_paragraph.paragraph_format.first_line_indent = Pt(50)

    # Figures
    doc.add_page_break()
    add_heading_with_color('FIGURES', level=1)
    sorted_claims = sorted(
        response["claims"], key=lambda x: x.get("claim_type") == "system")

    for claim in sorted_claims:
        if claim.get("generated_figures_data") and claim["generated_figures_data"].get("latex_details"):
            claim_type = claim.get("claim_type")
            for latex_detail in claim["generated_figures_data"]["latex_details"]:
                if "images_urls" in latex_detail:
                    for img_url in latex_detail["images_urls"]:
                        img_response = requests.get(img_url)
                        if img_response.status_code == 200:
                            image_stream = BytesIO(img_response.content)
                            image_paragraph = doc.add_paragraph()
                            run = image_paragraph.add_run()
                            run.add_picture(image_stream, width=Inches(4))
                            image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            print(
                                f"Failed to download the image from {img_url}")

    # Add "Attorney Docket No." to the header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_paragraph.add_run("Attorney Docket No.")

    # Function to add page numbering in footer
    def add_page_numbering(doc):
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.add_paragraph()
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = footer_paragraph.add_run()
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
        )
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE  \\* MERGEFORMAT </w:instrText>'
            )
        )
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
        )

    add_page_numbering(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


  