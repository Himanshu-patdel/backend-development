# from docx import Document
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.shared import Pt

# # Create a new Document object
# doc = Document()

# # Add the title with left alignment and simulate equal width
# title_text = 'JSON Response Handling Documentation'
# title_paragraph = doc.add_paragraph(title_text)
# title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# # Customize font size for the title (optional)
# run = title_paragraph.runs[0]
# run.font.size = Pt(24)

# # Introduction
# doc.add_heading('Introduction', level=1)
# doc.add_paragraph(
#     "JSON (JavaScript Object Notation) is a lightweight format for transmitting data between a client and a server. "
#     "In this project, we handle JSON responses from our APIs to perform various actions such as fetching data or processing requests."
# )

# # Prerequisites
# doc.add_heading('Prerequisites', level=1)
# doc.add_paragraph("Before starting, ensure you have the following:")
# doc.add_paragraph("- Python 3.x")
# doc.add_paragraph("- `requests` library for making API calls. You can install it using the following command:")
# doc.add_paragraph("```bash\npip install requests\n```")

# # Making API Calls
# doc.add_heading('Making API Calls', level=1)
# doc.add_paragraph("Here’s an example of how to send API requests and retrieve JSON responses using Python:")
# doc.add_paragraph(
#     """```python
# import requests

# # Example API call to get a JSON response
# response = requests.get('https://api.example.com/data')

# # Check if the request was successful
# if response.status_code == 200:
#     # Parse the JSON response
#     json_data = response.json()
#     print(json_data)
# else:
#     print(f"Request failed with status code: {response.status_code}")
# ```"""
# )

# # Parsing JSON Data
# doc.add_heading('Parsing JSON Data', level=1)
# doc.add_paragraph("Once the JSON response is received, you can parse it to access specific fields:")
# doc.add_paragraph(
#     """```python
# # Accessing specific fields from the JSON response
# user_name = json_data['user']['name']
# email = json_data['user']['email']
# print(f"User Name: {user_name}, Email: {email}")
# ```"""
# )

# # Handling Errors
# doc.add_heading('Handling Errors', level=1)
# doc.add_paragraph(
#     "When dealing with API calls, handling errors is important. "
#     "Here’s an example of how to handle HTTP errors and JSON parsing errors:"
# )
# doc.add_paragraph(
#     """```python
# try:
#     response = requests.get('https://api.example.com/data')
#     response.raise_for_status()  # Check for HTTP errors
#     json_data = response.json()  # Convert response to JSON
# except requests.exceptions.HTTPError as http_err:
#     print(f"HTTP error occurred: {http_err}")
# except ValueError as json_err:
#     print(f"JSON decoding failed: {json_err}")
# except Exception as err:
#     print(f"An error occurred: {err}")
# ```"""
# )

# # Working with Complex JSON
# doc.add_heading('Working with Complex JSON', level=1)
# doc.add_paragraph("If the JSON response has nested objects or arrays, you can loop through them like this:")
# doc.add_paragraph(
#     """```python
# # Assuming the JSON has nested arrays/objects
# for item in json_data['items']:
#     product_name = item['product']['name']
#     price = item['price']
#     print(f"Product: {product_name}, Price: {price}")
# ```"""
# )

# # Example Response Format
# doc.add_heading('Example Response Format', level=1)
# doc.add_paragraph("Here’s an example of a typical JSON response format:")
# doc.add_paragraph(
#     """```json
# {
#   "user": {
#     "id": 12345,
#     "name": "John Doe",
#     "email": "johndoe@example.com"
#   },
#   "items": [
#     {
#       "product": {
#         "id": 1,
#         "name": "Laptop"
#       },
#       "price": 1200.00
#     },
#     {
#       "product": {
#         "id": 2,
#         "name": "Phone"
#       },
#       "price": 800.00
#     }
#   ]
# }
# ```"""
# )

# # Testing JSON Response (Optional)
# doc.add_heading('Testing JSON Response (Optional)', level=1)
# doc.add_paragraph(
#     "You can use tools like Postman to test JSON responses. "
#     "Here’s how to do it with Postman:"
# )
# doc.add_paragraph(
#     "1. Open Postman and select the request type (GET, POST, etc.).\n"
#     "2. Enter the API URL.\n"
#     "3. Click **Send** and observe the response in the JSON tab.\n"
#     "4. Use Postman’s code generation feature to create code snippets for handling the response."
# )

# # Conclusion
# doc.add_heading('Conclusion', level=1)
# doc.add_paragraph(
#     "JSON responses are a critical part of interacting with APIs. "
#     "By following these steps, you can efficiently handle and parse JSON data in your applications, allowing for smoother integration and data processing."
# )

# # Save the document
# doc.save('json_response.docx')

# print("Documentation has been generated as 'json_response_handling.docx'.")


from docx import Document
from docx.shared import Inches, RGBColor, Pt  # Import Pt for font size
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import requests
import docx  # Import docx for XML manipulation

# Create a new Document
doc = Document()
base_url = "http://192.168.29.6:8000/api"

def generate_response_to_doc():
    url = f"{base_url}/v1/drafting/history_report?search_id=d4ccf0ca-807c-11ef-8c5d-0242ac120002"
    response = requests.get(url)
    return response.json()

response = generate_response_to_doc()

def convert_json_to_doc_buffer(response)->str:
    """AI is creating summary for convert_json_to_doc_buffer

    Args:
        response ([type]): [description]

    Returns:
        str: [description]
    """

    def add_heading_with_color(text, level):
        """AI is creating summary for add_heading_with_color

        Args:
            text ([type]): [description]
            level ([type]): [description]
        """
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set color to black
            run.font.name = 'Times New Roman'  # Set font to Times New Roman
            run.font.size = Pt(12)  # Set font size to 12
        doc.add_paragraph() 

    # Title
    title_heading = doc.add_heading(response["title"]["text"].upper(), level=1)
    title_run = title_heading.runs[0]
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
    counter = 1

    # Background
    add_heading_with_color('BACKGROUND', level=1)

    # Add technical field text with numbering at the start
    technical_field = response["technical_field"]["text"]
    technical_paragraph = doc.add_paragraph()
    technical_run = technical_paragraph.add_run(f"[{counter:04d}] ")
    technical_run.bold = True
    technical_run.font.name = 'Times New Roman'
    technical_run.font.size = Pt(12)
    tech_field_run = technical_paragraph.add_run(technical_field)
    tech_field_run.font.name = 'Times New Roman'
    tech_field_run.font.size = Pt(12)
    technical_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1  

    background_text = response["background"]["text"]
    sections = background_text.split('\n\n')
    for section in sections:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[{counter:04d}] ")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        section_run = paragraph.add_run(section)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        counter += 1   

    add_heading_with_color('BRIEF SUMMARY', level=1)  

    summary_field = response["summary"]["text"]
    summary_paragraph = doc.add_paragraph()
    summary_run = summary_paragraph.add_run(f"[{counter:04d}] ")
    summary_run.bold = True
    summary_run.font.name = 'Times New Roman'
    summary_run.font.size = Pt(12)
    summary_field_run = summary_paragraph.add_run(summary_field)
    summary_field_run.font.name = 'Times New Roman'
    summary_field_run.font.size = Pt(12)
    summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1  

    # List of figures 
    add_heading_with_color('BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS', level=1)
    for figure in response["list_of_figures"]:
        figure_paragraph = doc.add_paragraph()
        figure_run = figure_paragraph.add_run(f"[{counter:04d}] ")
        figure_run.bold = True
        figure_run.font.name = 'Times New Roman'
        figure_run.font.size = Pt(12)
        figure_text_run = figure_paragraph.add_run(figure)
        figure_text_run.font.name = 'Times New Roman'
        figure_text_run.font.size = Pt(12)
        figure_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        counter += 1  

    add_heading_with_color('DETAILED DESCRIPTION', level=1)
    method_desc_text = response["description"]["method_desc"]["text"]
    method_desc_sections = method_desc_text.split('\n\n')

    for section in method_desc_sections:
        method_desc_paragraph = doc.add_paragraph()
        method_desc_run = method_desc_paragraph.add_run(f"[{counter:04d}] ")
        method_desc_run.bold = True
        method_desc_run.font.name = 'Times New Roman'
        method_desc_run.font.size = Pt(12)
        section_run = method_desc_paragraph.add_run(section)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        method_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        counter += 1   

    invention_desc_text = response["description"]["invention_desc"]["text"]
    invention_desc_sections = invention_desc_text.split('\n\n')

    for section in invention_desc_sections:
        invention_desc_paragraph = doc.add_paragraph()
        invention_desc_run = invention_desc_paragraph.add_run(f"[{counter:04d}] ")
        invention_desc_run.bold = True
        invention_desc_run.font.name = 'Times New Roman'
        invention_desc_run.font.size = Pt(12)
        section_run = invention_desc_paragraph.add_run(section)
        section_run.font.name = 'Times New Roman'
        section_run.font.size = Pt(12)
        invention_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        counter += 1  

    doc.add_page_break() 
    claims_heading = doc.add_heading('CLAIMS', level=1)
    claims_run = claims_heading.runs[0]
    claims_run.font.color.rgb = RGBColor(0, 0, 0)   
    claims_run.font.name = 'Times New Roman'
    claims_run.font.size = Pt(12)
    claims_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  

    # Loop through claims and add numbering
    for index, claim in enumerate(response['claims'], start=1):
        claim_paragraph = doc.add_paragraph()
        claim_run = claim_paragraph.add_run(f"{index}. ")
        claim_run.font.name = 'Times New Roman'
        claim_run.font.size = Pt(12)
        claim_text_run = claim_paragraph.add_run(claim['text'])
        claim_text_run.font.name = 'Times New Roman'
        claim_text_run.font.size = Pt(12)
        claim_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Abstract
    doc.add_page_break() 
    abstract_heading = doc.add_heading('ABSTRACT', level=1)
    abstract_run = abstract_heading.runs[0]
    abstract_run.font.color.rgb = RGBColor(0, 0, 0)   
    abstract_run.font.name = 'Times New Roman'
    abstract_run.font.size = Pt(12)
    abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  

    abstract_paragraph = doc.add_paragraph(response["abstract"]["text"])
    # Ensure all runs in the abstract paragraph are Times New Roman
    for run in abstract_paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    abstract_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Figures
    doc.add_page_break() 
    add_heading_with_color('FIGURES', level=1)

    # Loop through claims and extract images
    for claim in response["claims"]:
        if claim.get("generated_figures_data") and claim["generated_figures_data"].get("latex_details"):
            for latex_detail in claim["generated_figures_data"]["latex_details"]:
                if "images_urls" in latex_detail:
                    for img_url in latex_detail["images_urls"]:
                        img_response = requests.get(img_url)
                        
                        if img_response.status_code == 200:
                            image_stream = BytesIO(img_response.content)
                            image_paragraph = doc.add_paragraph()
                            run = image_paragraph.add_run()
                            run.add_picture(image_stream, width=Inches(4))
                            image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            print(f"Failed to download the image from {img_url}")

    # Add "Attorney Docket No." to the header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_paragraph.add_run("Attorney Docket No.")

    # Function to add page numbering
    # Function to add page numbering
    def add_page_numbering(doc):
        section = doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.add_paragraph()
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # footer_paragraph.add_run("Page ")

        # Add page number field with correct namespace
        run = footer_paragraph.add_run()
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
        )
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE  \\* MERGEFORMAT </w:instrText>'
            )
        )
        run._element.append(
            docx.oxml.parse_xml(
                r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
        )


    # Call the function to add page numbering
    add_page_numbering(doc)

    # Save the document
    doc.save('testing1234.docx')
    print("Drafting document saved as testing.docx")
      # Save the document to a BytesIO buffer instead of a file
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Move to the beginning of the buffer

    return buffer




#   page_number_run = footer_paragraph.add_run()
#     page_number_run._element.append(
#         parse_xml(
#             r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
#         )
#     )
#     page_number_run._element.append(
#         parse_xml(
#             r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE  \\* MERGEFORMAT </w:instrText>'
#         )
#     )
#     page_number_run._element.append(
#         parse_xml(
#             r'<w:fldChar w:fldCharType="separate" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
#         )
#     )
#     page_number_run._element.append(
#         parse_xml(
#             r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
#         )
#     )





# page_number_run._element.append(
#     parse_xml(
#         r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
#     )
# )
# page_number_run._element.append(
#     parse_xml(
#         r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE  \\* MERGEFORMAT </w:instrText>'
#     )
# )
# page_number_run._element.append(
#     parse_xml(
#         r'<w:fldChar w:fldCharType="separate" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
#     )
# )
# page_number_run._element.append(
#     parse_xml(
#         r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
#     )
# )
