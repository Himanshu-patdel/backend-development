# def justify_text(text, line_length):
#     words = text.split()
#     lines = []
#     current_line = []

#     # Build lines based on the specified line length
#     for word in words:
#         if sum(len(w) for w in current_line) + len(current_line) + len(word) > line_length:
#             lines.append(current_line)
#             current_line = [word]
#         else:
#             current_line.append(word)
    
#     if current_line:
#         lines.append(current_line)

#     # Justify each line
#     justified_lines = []
#     for line in lines:
#         if len(line) == 1:
#             justified_lines.append(line[0])
#         else:
#             total_chars = sum(len(word) for word in line)
#             total_spaces = line_length - total_chars
#             space_between_words = total_spaces // (len(line) - 1)
#             extra_spaces = total_spaces % (len(line) - 1)

#             justified_line = ""
#             for i in range(len(line) - 1):
#                 justified_line += line[i] + ' ' * (space_between_words + (1 if i < extra_spaces else 0))
#             justified_line += line[-1]  # add the last word
#             justified_lines.append(justified_line)

#     return '\n'.join(justified_lines)

# # Input text and desired line length
# # text = ("Dynamic allocation of IP addresses based on network demand, "
# #         "utilizing the IPv6 protocol for expanded address space, "
# #         "implementing IP address translation for network security, "
# #         "integrating an IP address management system for efficient resource allocation, "
# #         "introducing IP address anonymization techniques for privacy protection, "
# #         "and developing IP address geolocation technology for targeted content delivery.")
# text='''
# Dynamic allocation of IP addresses based on network demand, utilizing the IPv6 protocol for expanded address space, implementing IP address translation for network security, integrating an IP address management system for efficient resource allocation, introducing IP address anonymization techniques for privacy protection, and developing IP address geolocation technology for targeted content delivery.
 
# [0013] Step 102 focuses on the integration of an IP address management system within the network infrastructure. This integration is essential for the efficient allocation of network resources. The IP address management system helps in organizing and managing IP addresses effectively, ensuring that resources are allocated optimally based on the network's requirements. This step streamlines the allocation process and helps in preventing IP address conflicts or wastage.

# '''
# line_length = 100  # Adjust this as needed

# # Generate and print the justified text
# justified_paragraph = justify_text(text, line_length)
# print("\nJustified Paragraph:\n")
# print(justified_paragraph)


from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import requests
from docx.oxml import OxmlElement

doc = Document()
base_url = "http://192.168.0.103:8000/api"

def generate_response_to_doc():
    url = f"{base_url}/v1/drafting/history_report?search_id=c81598b2-7fb6-11ef-917f-0242ac120002"
    response = requests.get(url)
    return response.json()

response = generate_response_to_doc()

# Function to add headings with black color
def add_heading_with_color(text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set color to black

# Add the title with black color
add_heading_with_color(response["title"]["text"], level=0)

# Background
add_heading_with_color('Background', level=1)
bg_paragraph = doc.add_paragraph(response["background"]["text"])
bg_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Brief Summary
add_heading_with_color('BRIEF SUMMARY', level=1)
summary_paragraph = doc.add_paragraph(response["summary"]["text"])
summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Brief Description of the Several Views of the Drawings
add_heading_with_color('BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS', level=1)
desc_paragraph = doc.add_paragraph(response["list_of_figures"])
desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Detailed Description
add_heading_with_color('DETAILED DESCRIPTION', level=1)
method_desc_paragraph = doc.add_paragraph(response["description"]["method_desc"]["text"])
method_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

invention_desc_paragraph = doc.add_paragraph(response["description"]["invention_desc"]["text"])
invention_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Claims
add_heading_with_color('CLAIMS', level=1)
for claim in response['claims']:
    claim_paragraph = doc.add_paragraph(claim['text'])
    claim_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    doc.add_paragraph()  # for space

# Abstract
add_heading_with_color('ABSTRACT', level=1)
abstract_paragraph = doc.add_paragraph(response["abstract"]["text"])
abstract_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Figures
add_heading_with_color('FIGURES', level=1)

# Loop through claims and extract images
for claim in response["claims"]:
    if claim.get("generated_figures_data") and claim["generated_figures_data"].get("latex_details"):
        for latex_detail in claim["generated_figures_data"]["latex_details"]:
            if "images_urls" in latex_detail:
                for img_url in latex_detail["images_urls"]:
                    # Download the image
                    img_response = requests.get(img_url)
                    
                    if img_response.status_code == 200:
                        image_stream = BytesIO(img_response.content)
                        doc.add_picture(image_stream, width=Inches(4))
                        doc.add_paragraph(latex_detail["fig_name"])
                    else:
                        print(f"Failed to download the image from {img_url}")

# Add "Attorney Docket No." to the header
section = doc.sections[0]
header = section.header
header_paragraph = header.add_paragraph()
header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
header_paragraph.add_run("Attorney Docket No.")

# Add page numbers to the footer
footer = section.footer
footer_paragraph = footer.add_paragraph()
footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add page number field using XML
field_code = OxmlElement('w:fldChar')
field_code.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
footer_paragraph._element.append(field_code)

instr_text = OxmlElement('w:instrText')
instr_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
instr_text.text = ' PAGE '
footer_paragraph._element.append(instr_text)

field_code = OxmlElement('w:fldChar')
field_code.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'separate')
footer_paragraph._element.append(field_code)

number_run = footer_paragraph.add_run()
number_run.text = "1"  # Placeholder; will show actual page number in Word

field_code = OxmlElement('w:fldChar')
field_code.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
footer_paragraph._element.append(field_code)

# Save the document
doc.save('drafting_check.docx')
print("Drafting document saved as drafting.docx")










# Initialize the counter
counter = 1

# Detailed Description Section
add_heading_with_color('DETAILED DESCRIPTION', level=1)
method_desc_text = response["description"]["method_desc"]["text"]
method_desc_sections = method_desc_text.split('\n\n')

# Adding Method Description
for section in method_desc_sections:
    method_desc_paragraph = doc.add_paragraph()
    method_desc_paragraph.add_run(f"[{counter:04d}] ").bold = True  
    method_desc_paragraph.add_run(section.strip())  # Strip whitespace
    method_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1   

# System Description Section
add_heading_with_color('SYSTEM DESCRIPTION', level=1)

# Loop through the text_list and add numbering for each item
for index, item in enumerate(response["description"]["system_desc"]["text_list"], start=counter):
    system_desc_paragraph = doc.add_paragraph()
    system_desc_paragraph.add_run(f"[{index:04d}] ").bold = True  # Use the current index
    system_desc_paragraph.add_run(item.strip())  # Adding the item text, stripped of leading/trailing whitespace
    system_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Update counter to the next value
counter += len(response["description"]["system_desc"]["text_list"])

# Invention Description Section
add_heading_with_color('INVENTION DESCRIPTION', level=1)
invention_desc_text = response["description"]["invention_desc"]["text"]
invention_desc_sections = invention_desc_text.split('\n\n')

# Adding Invention Description
for section in invention_desc_sections:
    invention_desc_paragraph = doc.add_paragraph()
    invention_desc_paragraph.add_run(f"[{counter:04d}] ").bold = True   
    invention_desc_paragraph.add_run(section.strip())  # Strip whitespace
    invention_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1  
