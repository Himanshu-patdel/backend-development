from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import requests

# Load the template
# doc = Document("my-template.docx")  # Use your template with page numbers
doc = Document("/Users/patdelanalytics/backend-development/hello.docx")

base_url = "http://192.168.29.6:8000/api"

def generate_response_to_doc():
    url = f"{base_url}/v1/drafting/history_report?search_id=701cdca8-7fff-11ef-8b42-0242ac120002"
    response = requests.get(url)
    return response.json()

response = generate_response_to_doc()

def add_heading_with_color(text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set color to black

# Title
title_heading = doc.add_heading(response["title"]["text"].upper(), level=1)
title_run = title_heading.runs[0]
title_run.font.color.rgb = RGBColor(0, 0, 0)
title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
counter = 1

# Background Section
add_heading_with_color('BACKGROUND', level=1)
technical_field = response["technical_field"]["text"]
technical_paragraph = doc.add_paragraph()
technical_paragraph.add_run(f"[{counter:04d}] ").bold = True
technical_paragraph.add_run(technical_field)
technical_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
counter += 1  

background_text = response["background"]["text"]
sections = background_text.split('\n\n')
for section in sections:
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"[{counter:04d}] ").bold = True  
    paragraph.add_run(section)  
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1   

# Brief Summary Section
add_heading_with_color('BRIEF SUMMARY', level=1)  
summary_field = response["summary"]["text"]
summary_paragraph = doc.add_paragraph()
summary_paragraph.add_run(f"[{counter:04d}] ").bold = True
summary_paragraph.add_run(summary_field)
summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
counter += 1  

# List of Figures
add_heading_with_color('BRIEF DESCRIPTION OF THE SEVERAL VIEWS OF THE DRAWINGS', level=1)
for figure in response["list_of_figures"]:
    figure_paragraph = doc.add_paragraph()
    figure_paragraph.add_run(f"[{counter:04d}] ").bold = True   
    figure_paragraph.add_run(figure)
    figure_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1  

# Detailed Description
add_heading_with_color('DETAILED DESCRIPTION', level=1)
method_desc_text = response["description"]["method_desc"]["text"]
method_desc_sections = method_desc_text.split('\n\n')

for section in method_desc_sections:
    method_desc_paragraph = doc.add_paragraph()
    method_desc_paragraph.add_run(f"[{counter:04d}] ").bold = True  
    method_desc_paragraph.add_run(section)
    method_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1   

# Invention Description
invention_desc_text = response["description"]["invention_desc"]["text"]
invention_desc_sections = invention_desc_text.split('\n\n')

for section in invention_desc_sections:
    invention_desc_paragraph = doc.add_paragraph()
    invention_desc_paragraph.add_run(f"[{counter:04d}] ").bold = True   
    invention_desc_paragraph.add_run(section)
    invention_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    counter += 1  

# Claims Section
doc.add_page_break() 
claims_heading = doc.add_heading('CLAIMS', level=1)
title_run = claims_heading.runs[0]
title_run.font.color.rgb = RGBColor(0, 0, 0)   
claims_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
for index, claim in enumerate(response['claims'], start=1):
    claim_paragraph = doc.add_paragraph()
    claim_paragraph.add_run(f"{index}. ")
    claim_paragraph.add_run(claim['text'])  
    claim_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Abstract Section
doc.add_page_break() 
abstract_heading = doc.add_heading('ABSTRACT', level=1)
title_run = abstract_heading.runs[0]
title_run.font.color.rgb = RGBColor(0, 0, 0)   
abstract_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
abstract_paragraph = doc.add_paragraph(response["abstract"]["text"])
abstract_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Figures Section
doc.add_page_break() 
add_heading_with_color('FIGURES', level=1)
for claim in response["claims"]:
    if claim.get("generated_figures_data") and claim["generated_figures_data"].get("latex_details"):
        for latex_detail in claim["generated_figures_data"]["latex_details"]:
            if "images_urls" in latex_detail:
                for img_url in latex_detail["images_urls"]:
                    img_response = requests.get(img_url)
                    if img_response.status_code == 200:
                        image_stream = BytesIO(img_response.content)
                        image_paragraph = doc.add_paragraph()
                        run = image_paragraph.add_run()
                        run.add_picture(image_stream, width=Inches(4))
                        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  
                    else:
                        print(f"Failed to download the image from {img_url}")

# Save the document
doc.save('drafting_ans_original1.docx')
print("Drafting document saved as drafting_ans_original.docx")
